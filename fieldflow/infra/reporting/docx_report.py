from __future__ import annotations

from pathlib import Path
from datetime import datetime

from docx import Document

from fieldflow.app.compare import ChangeSummary


def export_impact_report_docx(
    path: str | Path,
    scenario_name: str,
    baseline_duration: int | None,
    scenario_duration: int | None,
    changes: ChangeSummary,
) -> Path:
    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("FieldFlow Impact Report", level=1)
    doc.add_paragraph(f"Scenario: {scenario_name}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(f"Baseline project duration: {baseline_duration if baseline_duration is not None else '—'} days")
    doc.add_paragraph(f"Scenario project duration: {scenario_duration if scenario_duration is not None else '—'} days")

    if baseline_duration is not None and scenario_duration is not None:
        delta = scenario_duration - baseline_duration
        doc.add_paragraph(f"Delta (scenario - baseline): {delta} days")

    doc.add_heading("Changes", level=2)

    doc.add_heading("Duration changes", level=3)
    if changes.changed_durations:
        table = doc.add_table(rows=1, cols=3)
        hdr = table.rows[0].cells
        hdr[0].text = "Activity ID"
        hdr[1].text = "Baseline (days)"
        hdr[2].text = "Scenario (days)"
        for aid, bd, sd in changes.changed_durations:
            row = table.add_row().cells
            row[0].text = aid
            row[1].text = str(bd)
            row[2].text = str(sd)
    else:
        doc.add_paragraph("(none)")

    doc.add_heading("Relationships added", level=3)
    if changes.added_relationships:
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "Pred"
        hdr[1].text = "Succ"
        hdr[2].text = "Type"
        hdr[3].text = "Lag"
        for r in changes.added_relationships:
            row = table.add_row().cells
            row[0].text = r.pred_id
            row[1].text = r.succ_id
            row[2].text = r.rel_type.value
            row[3].text = str(r.lag_days)
    else:
        doc.add_paragraph("(none)")

    doc.add_heading("Relationships removed", level=3)
    if changes.removed_relationships:
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "Pred"
        hdr[1].text = "Succ"
        hdr[2].text = "Type"
        hdr[3].text = "Lag"
        for r in changes.removed_relationships:
            row = table.add_row().cells
            row[0].text = r.pred_id
            row[1].text = r.succ_id
            row[2].text = r.rel_type.value
            row[3].text = str(r.lag_days)
    else:
        doc.add_paragraph("(none)")

    doc.add_heading("Assumptions (v0)", level=2)
    doc.add_paragraph("This report reflects CPM calculations using activity durations, relationship types, and lags.")
    doc.add_paragraph("Calendars, constraints, and resource leveling are not included in this version.")

    doc.save(str(p))
    return p